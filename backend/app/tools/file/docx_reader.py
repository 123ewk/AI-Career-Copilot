"""DOCX Reader 异步文本提取工具

职责:
- 提供 Word(docx) 文件的异步文本提取能力
- 支持文件路径与字节流两种入参方式
- 输出结构化结果(纯文本、段落数、表格文本)

设计动机:
- 与 PDF Reader 配套,共同构成 Resume Parser 的输入管道
- python-docx 是 Python 生态读取 docx 的事实标准,直接使用
- 第三方库均为同步实现,使用 asyncio.to_thread 包装到线程池,
  避免阻塞 Event Loop

并发模型:
- 与 PDFReader 相同:IO 与解析都丢到默认 ThreadPoolExecutor
- 多个 docx 并发解析会被分配到不同线程,互不阻塞

异常映射:
- 文件不存在 / 输入为空 / 超过大小限制 -> ValidationError(400)
- docx 解析失败(非 zip/结构损坏) -> ExternalServiceError(502)
"""

from __future__ import annotations

import asyncio
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document as _DocxDocument
from docx.document import Document as _DocumentType
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import (
    ExternalServiceError,
    ValidationError,
)
from app.core.logger import logger
from app.core.settings import get_settings

# ==================== 数据结构 ====================

@dataclass(slots=True)
class DOCXReadResult:
    """DOCX 读取结果

    Attributes:
        text: 全部正文 + 表格的纯文本(段落间换行,表格行内 tab 分隔)
        paragraph_count: 正文段落数(不含空段)
        table_count: 表格数量
        core_properties: 文档核心属性(title/author/created 等),无则为空 dict
    """

    text: str
    paragraph_count: int
    table_count: int
    core_properties: dict[str, str] = field(default_factory=dict)


# ==================== 内部辅助 ====================

def _resolve_size_limit() -> int:
    """从 settings 读取最大允许字节数

    复用全局上传大小限制,避免 Reader 维护独立的配置。
    """
    return get_settings().max_upload_size_mb * 1024 * 1024


def _ensure_bytes_size(content: bytes) -> None:
    """校验字节流大小

    与 PDF Reader 同等防御策略:
    - 解析是 CPU + IO 密集型,大文件会长时间占用线程
    - 二次防御上传层遗漏的异常大文件
    """
    if not content:
        raise ValidationError(
            detail="DOCX 内容为空",
            error_code="VAL_010",
            extra={"size": 0},
        )
    size_limit = _resolve_size_limit()
    if len(content) > size_limit:
        raise ValidationError(
            detail=f"DOCX 文件超过大小限制({size_limit // (1024 * 1024)}MB)",
            error_code="VAL_011",
            extra={"size": len(content), "limit": size_limit},
        )


def _read_path_sync(path: Path) -> bytes:
    """同步读取本地文件,在线程池内执行

    用 'rb' 模式保留原始字节,后续统一走 bytes 解析路径,减少分支。
    """
    if not path.exists():
        raise ValidationError(
            detail=f"DOCX 文件不存在: {path}",
            error_code="VAL_012",
            extra={"path": str(path)},
        )
    if not path.is_file():
        raise ValidationError(
            detail=f"路径不是文件: {path}",
            error_code="VAL_013",
            extra={"path": str(path)},
        )
    return path.read_bytes()


def _extract_paragraphs_text(doc: _DocumentType) -> list[str]:
    """抽取所有正文章段的非空文本

    为何过滤空段:
    - python-docx 会把 Word 里的空行也识别为段落
    - 简历正文关注实际文字,空段会在后续 RAG chunking 时引入噪声
    """
    return [p.text for p in doc.paragraphs if p.text and p.text.strip()]


def _extract_tables_text(doc: _DocumentType) -> list[str]:
    """抽取所有表格文本,行内用 tab 分隔,行间用换行分隔

    简历里常见「项目经历|技能矩阵」等表格,需要纳入正文上下文。
    """
    table_chunks: list[str] = []
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            # 每行用 tab 拼接,避免「张三|开发|北京」糊成一团
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append("\t".join(cells))
        block = "\n".join(rows_text).strip()
        if block:
            table_chunks.append(block)
    return table_chunks


def _parse_docx_sync(content: bytes) -> DOCXReadResult:
    """同步解析 DOCX,纯 CPU 工作,放线程池避免阻塞

    失败模式:
    - BadZipFile / PackageNotFoundError -> 不是合法 docx(格式错误)
    - 其他 Exception -> 兜底为外部服务错误(文件损坏)
    """
    try:
        doc = _DocxDocument(BytesIO(content))
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        raise ExternalServiceError(
            detail="DOCX 文件无法解析,不是合法的 Word 文档",
            error_code="EXT_003",
            extra={"exception": type(exc).__name__},
        ) from exc
    except Exception as exc:
        # 兜底:docx 解析过程中可能因 XML 损坏抛 ValueError/KeyError 等
        raise ExternalServiceError(
            detail="DOCX 文件解析失败,文件可能已损坏",
            error_code="EXT_004",
            extra={"exception": type(exc).__name__},
        ) from exc

    # ---- 文本组装:段落在前,表格在后,之间用空行分隔 ----
    paragraph_chunks = _extract_paragraphs_text(doc)
    table_chunks = _extract_tables_text(doc)
    text = "\n".join(paragraph_chunks + table_chunks)

    # ---- 核心属性 ----
    # core_properties 是 docx 的 core.xml,字段可能为 None
    cp = doc.core_properties
    core_properties: dict[str, str] = {}
    for attr in ("title", "author", "subject", "keywords", "comments"):
        value = getattr(cp, attr, None)
        if value:
            core_properties[attr] = str(value)
    if cp.created:
        core_properties["created"] = cp.created.isoformat()
    if cp.modified:
        core_properties["modified"] = cp.modified.isoformat()

    return DOCXReadResult(
        text=text,
        paragraph_count=len(paragraph_chunks),
        table_count=len(doc.tables),
        core_properties=core_properties,
    )


# ==================== 公共 API ====================

class DOCXReader:
    """DOCX 异步读取器

    用法:
        reader = DOCXReader()
        result = await reader.read("resume.docx")
        print(result.text, result.paragraph_count)

    设计为可复用实例:Reader 是无状态的,实例方法只是入口,
    实际解析在 _parse_docx_sync 内完成。
    """

    async def read(self, path: str | Path) -> DOCXReadResult:
        """从本地路径异步读取 DOCX

        Args:
            path: DOCX 文件绝对或相对路径

        Returns:
            DOCXReadResult: 包含文本、段落数、表格数、核心属性

        Raises:
            ValidationError: 文件不存在/不是文件/超过大小限制
            ExternalServiceError: DOCX 解析失败(文件损坏/格式错误)
        """
        path_obj = Path(path)
        logger.info("DOCX 读取开始 | path={}", path_obj)

        content = await asyncio.to_thread(_read_path_sync, path_obj)
        _ensure_bytes_size(content)
        result = await asyncio.to_thread(_parse_docx_sync, content)

        logger.info(
            "DOCX 读取完成 | path={} | paragraphs={} | tables={} | text_len={}",
            path_obj,
            result.paragraph_count,
            result.table_count,
            len(result.text),
        )
        return result

    async def read_bytes(self, content: bytes) -> DOCXReadResult:
        """从字节流异步读取 DOCX

        典型用法:HTTP 上传文件、对象存储下载等已加载到内存的场景。

        Args:
            content: DOCX 文件的原始字节流

        Returns:
            DOCXReadResult: 包含文本、段落数、表格数、核心属性

        Raises:
            ValidationError: 内容类型错误/为空/超过大小限制
            ExternalServiceError: DOCX 解析失败
        """
        if not isinstance(content, (bytes, bytearray)):
            raise ValidationError(
                detail="content 必须是 bytes 类型",
                error_code="VAL_014",
                extra={"type": type(content).__name__},
            )
        content_bytes = bytes(content)
        _ensure_bytes_size(content_bytes)
        logger.info("DOCX 读取开始(bytes) | size={}", len(content_bytes))

        result = await asyncio.to_thread(_parse_docx_sync, content_bytes)

        logger.info(
            "DOCX 读取完成(bytes) | paragraphs={} | tables={} | text_len={}",
            result.paragraph_count,
            result.table_count,
            len(result.text),
        )
        return result
