"""PDF / DOCX Reader 单元测试

覆盖场景:
- 正常路径:从 path / bytes 读取
- 边界条件:空文件、超大文件、非文件路径、目录路径
- 类型错误:bytes 参数传 str
- 文件损坏:随机二进制 / 非 zip 数据
- 并发安全:多协程并发读同一文件,互不干扰
"""

from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.core.exceptions import ExternalServiceError, ValidationError
from app.tools.file import DOCXReader, PDFReader

# ==================== Fixture ====================

def _build_pdf_bytes(
    title: str = "Test Resume",
    author: str = "Test Suite",
    text: str = "Hello PDF Resume\nJohn Doe | john.doe@example.com",
) -> bytes:
    """构造一个包含文字的 PDF,返回字节流

    使用 pypdf 底层 API + DictionaryObject 构建,
    避免引入 reportlab 等额外依赖。
    """
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    # 字体字典(必须用 DictionaryObject,裸 dict 会被 pypdf 拒绝)
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")

    fonts = DictionaryObject()
    fonts[NameObject("/F1")] = font

    resources = DictionaryObject()
    resources[NameObject("/Font")] = fonts
    page[NameObject("/Resources")] = resources

    # 把每行转成 PDF 文字绘制指令(Tj 操作符)
    line_ops = b"".join(
        f"({line}) Tj 0 -20 Td ".encode("latin-1") for line in text.split("\n")
    )
    content_bytes = b"BT /F1 14 Tf 50 750 Td " + line_ops + b"ET"
    stream = DecodedStreamObject()
    stream.set_data(content_bytes)

    # 把 stream 注册为间接对象,再把引用挂到 page
    content_ref = writer._add_object(stream)
    page[NameObject("/Contents")] = content_ref

    writer.add_metadata(
        {
            "/Title": title,
            "/Author": author,
        }
    )

    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


@pytest.fixture
def pdf_path(tmp_path: Path) -> Path:
    """生成一个包含两段文字的 PDF 文件,返回路径"""
    path = tmp_path / "resume.pdf"
    path.write_bytes(_build_pdf_bytes())
    return path


@pytest.fixture
def docx_path(tmp_path: Path) -> Path:
    """生成一个包含段落和表格的 DOCX 文件,返回路径"""
    doc = DocxDocument()
    doc.add_heading("Jane Smith Resume", level=1)
    doc.add_paragraph("Email: jane.smith@example.com")
    doc.add_paragraph("Summary: Backend engineer with 5 years experience.")
    # 空段落用于验证空段过滤
    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Years"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "5"

    doc.core_properties.title = "Test DOCX"
    doc.core_properties.author = "Test Suite"

    path = tmp_path / "resume.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def pdf_bytes(pdf_path: Path) -> bytes:
    """把 fixture 生成的 PDF 读成 bytes"""
    return pdf_path.read_bytes()


@pytest.fixture
def docx_bytes(docx_path: Path) -> bytes:
    """把 fixture 生成的 DOCX 读成 bytes"""
    return docx_path.read_bytes()


# ==================== PDF Reader ====================

class TestPDFReader:
    """PDFReader 单元测试"""

    @pytest.mark.asyncio
    async def test_read_from_path_success(self, pdf_path: Path) -> None:
        """正常路径:从本地文件读取 PDF,文本与元数据均应正确"""
        reader = PDFReader()
        result = await reader.read(pdf_path)

        assert result.page_count == 1
        assert "Hello PDF Resume" in result.text
        assert "john.doe@example.com" in result.text
        assert result.metadata.get("/Title") == "Test Resume"
        assert result.metadata.get("/Author") == "Test Suite"

    @pytest.mark.asyncio
    async def test_read_from_bytes_success(self, pdf_bytes: bytes) -> None:
        """正常路径:从 bytes 读取 PDF"""
        reader = PDFReader()
        result = await reader.read_bytes(pdf_bytes)

        assert result.page_count == 1
        assert "Hello PDF Resume" in result.text
        assert result.metadata.get("/Title") == "Test Resume"

    @pytest.mark.asyncio
    async def test_read_bytes_accepts_bytearray(self, pdf_bytes: bytes) -> None:
        """类型兼容:bytearray 也应被接受(走 bytes() 转换)"""
        reader = PDFReader()
        result = await reader.read_bytes(bytearray(pdf_bytes))

        assert result.page_count == 1
        assert "Hello PDF Resume" in result.text

    @pytest.mark.asyncio
    async def test_read_bytes_rejects_non_bytes(self) -> None:
        """类型校验:传 str 应抛 ValidationError"""
        reader = PDFReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes("not bytes")  # type: ignore[arg-type]
        assert exc_info.value.error_code == "VAL_008"

    @pytest.mark.asyncio
    async def test_read_bytes_empty_raises(self) -> None:
        """空字节流应抛 ValidationError(VAL_002)"""
        reader = PDFReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes(b"")
        assert exc_info.value.error_code == "VAL_002"

    @pytest.mark.asyncio
    async def test_read_nonexistent_path_raises(self, tmp_path: Path) -> None:
        """不存在的文件路径应抛 ValidationError(VAL_004)"""
        reader = PDFReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read(tmp_path / "ghost.pdf")
        assert exc_info.value.error_code == "VAL_004"

    @pytest.mark.asyncio
    async def test_read_directory_raises(self, tmp_path: Path) -> None:
        """目录路径应抛 ValidationError(VAL_005)"""
        reader = PDFReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read(tmp_path)  # tmp_path 本身是目录
        assert exc_info.value.error_code == "VAL_005"

    @pytest.mark.asyncio
    async def test_read_oversized_bytes_raises(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """超大文件应抛 ValidationError(VAL_003)

        直接 patch 模块级 _resolve_size_limit,避免污染全局 settings 单例。
        """
        monkeypatch.setattr(
            "app.tools.file.pdf_reader._resolve_size_limit", lambda: 0
        )

        reader = PDFReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes(b"%PDF-1.4 fake")
        assert exc_info.value.error_code == "VAL_003"

    @pytest.mark.asyncio
    async def test_read_corrupt_bytes_raises(self) -> None:
        """损坏的 PDF 应抛 ExternalServiceError(EXT_002)"""
        reader = PDFReader()
        # 构造一个非 PDF 格式的字节流
        with pytest.raises(ExternalServiceError) as exc_info:
            await reader.read_bytes(b"this is not a pdf, just plain text")
        assert exc_info.value.error_code == "EXT_002"

    @pytest.mark.asyncio
    async def test_concurrent_reads_isolated(self, pdf_path: Path) -> None:
        """并发安全:多协程并发读同一文件,结果一致且互不干扰"""
        reader = PDFReader()
        results = await asyncio.gather(*(reader.read(pdf_path) for _ in range(10)))
        assert len(results) == 10
        assert all(r.page_count == 1 for r in results)
        assert all("Hello PDF Resume" in r.text for r in results)


# ==================== DOCX Reader ====================

class TestDOCXReader:
    """DOCXReader 单元测试"""

    @pytest.mark.asyncio
    async def test_read_from_path_success(self, docx_path: Path) -> None:
        """正常路径:从本地文件读取 DOCX,段落/表格/属性均应正确"""
        reader = DOCXReader()
        result = await reader.read(docx_path)

        assert result.paragraph_count == 3  # 标题 + 2 段正文(空段被过滤)
        assert result.table_count == 1
        assert "Jane Smith Resume" in result.text
        assert "jane.smith@example.com" in result.text
        assert "Python" in result.text
        # 表格行内 tab 分隔:校验 Skill 列与 Years 列
        assert "Skill\tYears" in result.text
        assert "Python\t5" in result.text
        # 核心属性
        assert result.core_properties.get("title") == "Test DOCX"
        assert result.core_properties.get("author") == "Test Suite"

    @pytest.mark.asyncio
    async def test_read_from_bytes_success(self, docx_bytes: bytes) -> None:
        """正常路径:从 bytes 读取 DOCX"""
        reader = DOCXReader()
        result = await reader.read_bytes(docx_bytes)

        assert result.paragraph_count == 3
        assert result.table_count == 1
        assert "Jane Smith Resume" in result.text

    @pytest.mark.asyncio
    async def test_read_bytes_accepts_bytearray(self, docx_bytes: bytes) -> None:
        """类型兼容:bytearray 也应被接受"""
        reader = DOCXReader()
        result = await reader.read_bytes(bytearray(docx_bytes))

        assert result.paragraph_count == 3
        assert "Jane Smith Resume" in result.text

    @pytest.mark.asyncio
    async def test_read_bytes_rejects_non_bytes(self) -> None:
        """类型校验:传 str 应抛 ValidationError"""
        reader = DOCXReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes("not bytes")  # type: ignore[arg-type]
        assert exc_info.value.error_code == "VAL_014"

    @pytest.mark.asyncio
    async def test_read_bytes_empty_raises(self) -> None:
        """空字节流应抛 ValidationError(VAL_010)"""
        reader = DOCXReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes(b"")
        assert exc_info.value.error_code == "VAL_010"

    @pytest.mark.asyncio
    async def test_read_nonexistent_path_raises(self, tmp_path: Path) -> None:
        """不存在的文件路径应抛 ValidationError(VAL_012)"""
        reader = DOCXReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read(tmp_path / "ghost.docx")
        assert exc_info.value.error_code == "VAL_012"

    @pytest.mark.asyncio
    async def test_read_directory_raises(self, tmp_path: Path) -> None:
        """目录路径应抛 ValidationError(VAL_013)"""
        reader = DOCXReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read(tmp_path)
        assert exc_info.value.error_code == "VAL_013"

    @pytest.mark.asyncio
    async def test_read_oversized_bytes_raises(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """超大文件应抛 ValidationError(VAL_011)

        直接 patch 模块级 _resolve_size_limit,避免污染全局 settings 单例。
        """
        monkeypatch.setattr(
            "app.tools.file.docx_reader._resolve_size_limit", lambda: 0
        )

        reader = DOCXReader()
        with pytest.raises(ValidationError) as exc_info:
            await reader.read_bytes(b"PK\x03\x04 fake")
        assert exc_info.value.error_code == "VAL_011"

    @pytest.mark.asyncio
    async def test_read_corrupt_bytes_raises(self) -> None:
        """非 docx 字节流应抛 ExternalServiceError(EXT_003)"""
        reader = DOCXReader()
        with pytest.raises(ExternalServiceError) as exc_info:
            await reader.read_bytes(b"not a docx, just plain text content")
        assert exc_info.value.error_code == "EXT_003"

    @pytest.mark.asyncio
    async def test_read_minimal_docx(self) -> None:
        """最小 docx:无段落无表格,应返回空 text 与零计数"""
        doc = DocxDocument()
        buf = BytesIO()
        doc.save(buf)
        reader = DOCXReader()
        result = await reader.read_bytes(buf.getvalue())

        assert result.text == ""
        assert result.paragraph_count == 0
        assert result.table_count == 0

    @pytest.mark.asyncio
    async def test_concurrent_reads_isolated(self, docx_path: Path) -> None:
        """并发安全:多协程并发读同一文件,结果一致"""
        reader = DOCXReader()
        results = await asyncio.gather(*(reader.read(docx_path) for _ in range(10)))
        assert len(results) == 10
        assert all(r.paragraph_count == 3 for r in results)
        assert all(r.table_count == 1 for r in results)
